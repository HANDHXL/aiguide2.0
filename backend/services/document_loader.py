"""Load and parse documents from the knowledge base.

Supported formats:
  - .docx  (Word documents — paragraphs + tables)
  - .xlsx  (Excel spreadsheets — summary statistics)
  - .pdf   (PDF documents — full text extraction)
  - .txt   (Plain text files)
  - .md    (Markdown files)
  - .csv   (CSV spreadsheets — same treatment as Excel)
"""

from pathlib import Path
from typing import List, Dict
from docx import Document
import pandas as pd


# ================================================================
#  Format-specific loaders
# ================================================================

def load_docx(file_path: Path) -> List[Dict[str, str]]:
    """Load and parse a .docx file into structured documents."""
    doc = Document(str(file_path))
    docs = []

    # Extract paragraphs
    paragraphs = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            paragraphs.append(text)

    # Extract tables
    tables = []
    for table in doc.tables:
        rows = []
        if not table.rows:
            continue
        headers = [cell.text.strip() for cell in table.rows[0].cells]
        for row in table.rows[1:]:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                row_dict = dict(zip(headers, cells))
                rows.append(row_dict)
        if rows:
            tables.append(rows)

    # Flatten paragraphs into a single document chunk
    if paragraphs:
        full_text = "\n\n".join(paragraphs)
        docs.append({
            "content": full_text,
            "source": file_path.name,
            "type": "docx_paragraphs"
        })

    # Each table row becomes a document chunk
    for t_idx, table_rows in enumerate(tables):
        for r_idx, row in enumerate(table_rows):
            row_text = " | ".join(f"{k}: {v}" for k, v in row.items() if v)
            docs.append({
                "content": row_text,
                "source": file_path.name,
                "type": f"docx_table_{t_idx}_row_{r_idx}",
                "metadata": row
            })

    return docs


def load_pdf(file_path: Path) -> List[Dict[str, str]]:
    """Load and parse a PDF file. Extracts text from all pages."""
    try:
        from pypdf import PdfReader
    except ImportError:
        raise ImportError("pypdf is required for PDF support. Install with: pip install pypdf")

    reader = PdfReader(str(file_path))
    docs = []

    # Strategy 1: per-page chunks (good for large PDFs)
    pages_text = []
    for i, page in enumerate(reader.pages):
        text = page.extract_text()
        if text and text.strip():
            pages_text.append(f"[第{i+1}页]\n{text.strip()}")

    if not pages_text:
        return []

    # Group pages into chunks of ~3 pages to balance context size
    chunk_size = 3
    for chunk_idx in range(0, len(pages_text), chunk_size):
        chunk_pages = pages_text[chunk_idx:chunk_idx + chunk_size]
        content = "\n\n".join(chunk_pages)
        if content.strip():
            page_range = f"p{chunk_idx+1}-p{min(chunk_idx+chunk_size, len(pages_text))}"
            docs.append({
                "content": content,
                "source": file_path.name,
                "type": f"pdf_{page_range}"
            })

    return docs


def load_excel(file_path: Path) -> List[Dict[str, str]]:
    """Load structured data from Excel / CSV for analytics context."""
    if file_path.suffix.lower() == ".csv":
        df = pd.read_csv(str(file_path))
    else:
        df = pd.read_excel(str(file_path))

    summaries = []
    numeric_cols = df.select_dtypes(include=["number"]).columns

    summary_lines = [f"数据集概览：共{len(df)}条记录，包含{len(df.columns)}个字段。"]
    summary_lines.append(f"字段列表：{', '.join(str(c) for c in df.columns.tolist())}。")

    for col in numeric_cols:
        if col in ("tourist_id", "id", "user_id"):
            continue
        stats = df[col].describe()
        summary_lines.append(
            f"{col} 统计：均值{stats['mean']:.2f}，"
            f"中位数{stats['50%']:.2f}，"
            f"最小值{stats['min']:.2f}，最大值{stats['max']:.2f}。"
        )

    # Satisfaction distribution
    if "satisfaction" in df.columns:
        sat_dist = df["satisfaction"].value_counts().sort_index()
        dist_parts = [f"满意度{k}分: {v}人({v/len(df)*100:.1f}%)" for k, v in sat_dist.items()]
        summary_lines.append(f"满意度分布：{'；'.join(dist_parts)}。")

    # Attraction popularity
    if "attraction_name" in df.columns:
        top_attractions = df["attraction_name"].value_counts().head(10)
        attr_parts = [f"{attr}: {cnt}次" for attr, cnt in top_attractions.items()]
        summary_lines.append(f"热门景点TOP10：{'；'.join(attr_parts)}。")

    # For small CSVs (< 100 rows), also include raw rows as chunks
    extra_chunks = []
    if len(df) < 100:
        for idx, row in df.iterrows():
            row_text = " | ".join(f"{col}: {val}" for col, val in row.items()
                                  if pd.notna(val) and str(val).strip())
            if row_text.strip():
                extra_chunks.append({
                    "content": row_text,
                    "source": file_path.name,
                    "type": f"row_{idx}",
                })

    result = [{
        "content": "\n".join(summary_lines),
        "source": file_path.name,
        "type": "data_summary"
    }]
    result.extend(extra_chunks)
    return result


def load_txt(file_path: Path) -> List[Dict[str, str]]:
    """Load a plain text file."""
    encodings = ["utf-8", "gbk", "gb2312", "latin-1"]
    for enc in encodings:
        try:
            text = file_path.read_text(encoding=enc)
            if text.strip():
                return [{
                    "content": text.strip(),
                    "source": file_path.name,
                    "type": "txt"
                }]
        except (UnicodeDecodeError, UnicodeError):
            continue
    return []


def load_md(file_path: Path) -> List[Dict[str, str]]:
    """Load a Markdown file. Strips basic formatting but preserves structure."""
    encodings = ["utf-8", "gbk", "gb2312", "latin-1"]
    for enc in encodings:
        try:
            raw = file_path.read_text(encoding=enc)
            if raw.strip():
                # Light cleanup: remove code fences but keep content
                import re
                text = re.sub(r'```[^`]*```', '', raw, flags=re.DOTALL)
                text = re.sub(r'`([^`]+)`', r'\1', text)
                text = re.sub(r'#{1,6}\s+', '', text)  # headers
                text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)  # bold
                text = re.sub(r'\*(.+?)\*', r'\1', text)      # italic
                text = re.sub(r'\[(.+?)\]\(.+?\)', r'\1', text)  # links
                text = re.sub(r'!\[.*?\]\(.+?\)', '', text)      # images
                text = re.sub(r'^[-*+]\s+', '• ', text, flags=re.MULTILINE)  # list bullets
                text = re.sub(r'\n{3,}', '\n\n', text)
                return [{
                    "content": text.strip(),
                    "source": file_path.name,
                    "type": "markdown"
                }]
        except (UnicodeDecodeError, UnicodeError):
            continue
    return []


# ================================================================
#  Unified loader — routes by file extension
# ================================================================

# Map file extensions → loader functions
_LOADER_MAP = {
    ".docx": load_docx,
    ".doc": load_docx,
    ".pdf": load_pdf,
    ".xlsx": load_excel,
    ".xls": load_excel,
    ".csv": load_excel,
    ".txt": load_txt,
    ".md": load_md,
    ".markdown": load_md,
}

# Ignore these files (temp files, office lock files, etc.)
_IGNORE_PATTERNS = {"~$", ".tmp", ".lock", ".DS_Store", "Thumbs.db"}


def load_all_documents(kb_dir: Path) -> List[Dict[str, str]]:
    """Load all supported documents from the knowledge base directory."""
    all_docs = []
    skipped = []

    for file_path in sorted(kb_dir.iterdir()):
        if not file_path.is_file():
            continue
        # Skip temp / lock files
        if any(file_path.name.startswith(p) for p in _IGNORE_PATTERNS):
            skipped.append(file_path.name)
            continue
        if any(file_path.suffix.lower() == p for p in _IGNORE_PATTERNS):
            skipped.append(file_path.name)
            continue

        suffix = file_path.suffix.lower()
        loader = _LOADER_MAP.get(suffix)
        if loader:
            try:
                docs = loader(file_path)
                all_docs.extend(docs)
                print(f"  [OK] {file_path.name} → {len(docs)} chunk(s)")
            except (ImportError, ValueError, UnicodeDecodeError, OSError) as e:
                print(f"  [SKIP] {file_path.name}: {e}")
                skipped.append(file_path.name)
        else:
            skipped.append(file_path.name)

    if skipped:
        print(f"  Skipped {len(skipped)} unsupported file(s): {skipped}")
    print(f"\n  Total: {len(all_docs)} document chunks from {len(all_docs) and len(set(d['source'] for d in all_docs)) or 0} files")
    return all_docs
